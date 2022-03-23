import docx
import sys
s = sys.stdin.read()
document = docx.Document()
time = s.split('\n')[1].split()[1]
place = s.split('\n')[0].lower().strip()
d = s.split('\n')
for i in range(len(s.split('\n')[2:])):
    document.add_heading('Приглашение', level=0)
    document.add_paragraph(f'Здравствуйте, {d[2 + i]}!\n'
                           f'Приглашаем вас отметить праздник 8 марта {place} в {time}.')
    document.add_page_break()
document.save('invitations.docx')
