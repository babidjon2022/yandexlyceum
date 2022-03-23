import docx


def markdown_to_docx(text):
    text = text.splitlines()
    document = docx.Document()
    document.add_heading(text[0], level=0)
    text = text[1:]
    for i in text:
        if not i:
            document.add_paragraph()
            continue
        while '  ' in i:
            i = i.replace('  ', '\n')
        if i[:3] == i[-3:] and i[:3] in {'___', '***'}:
            document.add_paragraph().add_run(i[3:-3]).bold = True
            document.paragraphs[-1].runs[-1].italic = True
            continue
        if i[:2] == i[-2:] and i[:2] in {'__', '**'}:
            document.add_paragraph().add_run(i[2:-2]).bold = True
            continue
        if i[0] == i[-1] and i[0] in {'_', '*'}:
            document.add_paragraph().add_run(i[1:-1]).italic = True
            continue
        if i.split()[0][-1] == '.' and i.split()[0][:-1].isdigit(
        ) and document.paragraphs[-1].style != 'Last Number':
            document.add_paragraph(' '.join(i.split()[1:]).strip(),
                                   style='List Number')
            continue
        elif i.split()[0][-1] == '.' and i.split()[0][:-1].isdigit():
            document.paragraphs[-1].add_run(' '.join(i.split()[1:]).strip())
            continue
        if i.split(
        )[0] == '-' and document.paragraphs[-1].style != 'List Bullet':
            document.add_paragraph(i[1:].strip(), style='List Bullet')
            continue
        elif i.split()[0] == '-':
            document.paragraphs[-1].add_run(i[1:].strip())
            continue
        if set(i.split()[0]) == {'#'}:
            document.add_heading(' '.join(i.split()[1:]),
                                 level=len(i.split()[0]))
            continue
        else:
            document.add_paragraph(i)
    document.save('res.docx')
